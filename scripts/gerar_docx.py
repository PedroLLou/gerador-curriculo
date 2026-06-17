# -*- coding: utf-8 -*-
"""
Gera a versao ATS-limpa (.docx): coluna unica, headings padrao, sem cor/icone.
Generate the ATS-clean (.docx) version: single column, standard headings, no color/icons.

Requisito / Requirement: pip install python-docx

COMO USAR / HOW TO USE:
  1. Edite o bloco "CONTEUDO" abaixo com os dados do curriculo desta vaga.
     Edit the "CONTENT" block below with this job's resume data.
  2. Rode / Run:  python scripts/gerar_docx.py
  3. Saida / Output: curriculo-ats.docx (na pasta atual / in the current folder)

Dica: peca ao assistente de IA para preencher o bloco CONTEUDO a partir do curriculo.md.
Tip: ask the AI assistant to fill the CONTENT block from curriculo.md.
"""
from docx import Document
from docx.shared import Pt

OUT = "curriculo-ats.docx"

# ============================ CONTEUDO / CONTENT ============================
# Edite estes valores. Use listas de strings para bullets.
# Edit these values. Use lists of strings for bullets.
NOME = "{{NOME}}"
TITULO = "{{TITULO}}"
CONTATO_1 = "{{LOCAL}} | {{EMAIL}} | {{TELEFONE}}"
CONTATO_2 = "{{GITHUB}} | {{LINKEDIN}}"
RESUMO = "{{RESUMO}}"

HABILIDADES = [
    "Core: {{SKILLS_CORE}}",
    "Web: {{SKILLS_WEB}}",
    "Banco de dados / Databases: {{SKILLS_DB}}",
    "Ferramentas / Tools: {{SKILLS_TOOLS}}",
    "Idiomas / Languages: {{IDIOMAS}}",
]

# Cada experiencia: (cabecalho em negrito, [bullets]) / each: (bold header, [bullets])
EXPERIENCIAS = [
    ("{{EXP1_TITULO}} | {{EXP1_META}}", [
        "{{EXP1_B1}}",
        "{{EXP1_B2}}",
        "{{EXP1_B3}}",
    ]),
]

PROJETOS = [
    "{{PROJ1_NOME}}: {{PROJ1_DESC}} {{PROJ1_URL}}",
]

FORMACAO = "{{FORMACAO}} | {{FORMACAO_META}}"
CERTIFICACOES = "{{CERTIFICACOES}}"
# ===========================================================================


def main():
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Pt(36)
        s.left_margin = s.right_margin = Pt(40)
    base = doc.styles["Normal"]
    base.font.name = "Calibri"
    base.font.size = Pt(10.5)

    def heading(t):
        p = doc.add_paragraph()
        r = p.add_run(t.upper()); r.bold = True; r.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)

    def bullet(t):
        p = doc.add_paragraph(t, style="List Bullet")
        p.paragraph_format.space_after = Pt(1)

    def line(t, bold=False, size=10.5, after=2):
        p = doc.add_paragraph()
        r = p.add_run(t); r.bold = bold; r.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(after)

    # Cabecalho / Header
    line(NOME, bold=True, size=18, after=0)
    line(TITULO, size=11, after=2)
    line(CONTATO_1, size=9.5, after=0)
    line(CONTATO_2, size=9.5, after=4)

    # Resumo / Summary
    heading("Resumo / Summary")
    line(RESUMO)

    # Habilidades / Skills
    heading("Habilidades / Skills")
    for item in HABILIDADES:
        bullet(item)

    # Experiencia / Experience
    heading("Experiencia / Experience")
    for header, bullets in EXPERIENCIAS:
        line(header, bold=True, after=0)
        for b in bullets:
            bullet(b)

    # Projetos / Projects
    if PROJETOS:
        heading("Projetos / Projects")
        for p in PROJETOS:
            bullet(p)

    # Formacao / Education
    heading("Formacao / Education")
    line(FORMACAO)

    # Certificacoes / Certifications
    if CERTIFICACOES:
        heading("Certificacoes / Certifications")
        line(CERTIFICACOES)

    doc.save(OUT)
    print("OK docx:", OUT)


if __name__ == "__main__":
    main()
